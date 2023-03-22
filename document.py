# mypy: ignore-errors
import os
import re

from docx import Document
from PIL import Image


doc = Document("sample.docx")
docnp = Document("sample_no_pic.docx")

# Detect a picture (file has pic)
for pg in doc.paragraphs:
    for runn in pg.runs:
        if "pic:pic" in runn.element.xml:
            print("Found a pic")

# Detect a picture (file does not have pic)
for pg in docnp.paragraphs:
    for runn in pg.runs:
        if "pic:pic" in runn.element.xml:
            print("Found a pic")

# Get a picture
for pg in doc.paragraphs:
    for runn in pg.runs:
        if "pic:pic" in runn.element.xml:
            part = doc.part
            rid = re.search(r'r:embed="([^"]+)"', runn.element.xml).group(1)
            image_part = part.related_parts[rid]
            result_file = "result_file"

            binary = image_part._blob
            with open(result_file, "wb") as file_handler:
                file_handler.write(binary)

            Image.open(result_file).save("test_save.png")
            os.remove(result_file)
